"""
Input File Validation System
Validates all required input files and configurations at startup
"""

import os
import json
import yaml
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches


@dataclass
class ValidationResult:
    """Result of input validation"""
    is_valid: bool
    errors: List[str]
    warnings: List[str]
    info: List[str]

    def add_error(self, message: str):
        self.errors.append(message)
        self.is_valid = False

    def add_warning(self, message: str):
        self.warnings.append(message)

    def add_info(self, message: str):
        self.info.append(message)


class InputValidator:
    """Comprehensive input validation for course content generator"""

    def __init__(self, base_path: str = "."):
        self.base_path = Path(base_path)
        self.result = ValidationResult(is_valid=True, errors=[], warnings=[], info=[])

    def validate_all(self) -> ValidationResult:
        """Run all validation checks"""
        print("🔍 Validating input files and configuration...")

        # Reset result
        self.result = ValidationResult(is_valid=True, errors=[], warnings=[], info=[])

        # Core validation checks
        self._validate_directory_structure()
        self._validate_secrets_configuration()
        self._validate_input_files()
        self._validate_configuration_files()
        self._validate_dependencies()

        # Summary
        if self.result.is_valid:
            self.result.add_info("✅ All validation checks passed")
        else:
            self.result.add_info(f"❌ Validation failed with {len(self.result.errors)} errors")

        return self.result

    def _validate_directory_structure(self):
        """Validate required directory structure exists"""
        required_dirs = [
            "input",
            "config",
            "temporal_output",
            "weekly_content",
            "run_logs",
            "app",
            "app/tools",
            "app/agents",
            "app/models",
            "app/utils",
            "app/workflow"
        ]

        for dir_name in required_dirs:
            dir_path = self.base_path / dir_name
            if not dir_path.exists():
                self.result.add_error(f"Required directory missing: {dir_path}")
            elif not dir_path.is_dir():
                self.result.add_error(f"Path exists but is not a directory: {dir_path}")

        self.result.add_info(f"📁 Directory structure validation completed")

    def _validate_secrets_configuration(self):
        """Validate secrets and API key configuration"""
        secrets_file = self.base_path / ".secrets"

        if not secrets_file.exists():
            self.result.add_warning(".secrets file not found - using environment variables only")
            self.result.add_info("💡 Copy .secrets.example to .secrets and configure your API keys")

        # Check for LLM configuration (either Azure OpenAI or regular OpenAI)
        azure_vars = ["AZURE_ENDPOINT", "AZURE_SUBSCRIPTION_KEY", "AZURE_API_VERSION"]
        azure_configured = all(os.getenv(var) for var in azure_vars)
        openai_configured = os.getenv("OPENAI_API_KEY") is not None

        if not (azure_configured or openai_configured):
            self.result.add_error("No LLM configuration found")
            self.result.add_info("Configure either Azure OpenAI or OpenAI API keys")
        elif azure_configured:
            self.result.add_info("🔑 Azure OpenAI configuration detected")
            # Validate Azure configuration
            endpoint = os.getenv("AZURE_ENDPOINT", "")
            if not endpoint.startswith("https://"):
                self.result.add_error("AZURE_ENDPOINT must start with https://")

            key = os.getenv("AZURE_SUBSCRIPTION_KEY", "")
            if len(key) < 32:  # Azure keys are typically much longer
                self.result.add_warning("AZURE_SUBSCRIPTION_KEY seems too short")

        elif openai_configured:
            self.result.add_info("🔑 OpenAI configuration detected")
            key = os.getenv("OPENAI_API_KEY", "")
            if not key.startswith("sk-"):
                self.result.add_warning("OPENAI_API_KEY should start with 'sk-'")

        # Check for web search provider
        search_providers = [
            "TAVILY_API_KEY",
            "BING_SEARCH_API_KEY",
            "SERPAPI_API_KEY"
        ]
        google_cse_configured = os.getenv("GOOGLE_CSE_KEY") and os.getenv("GOOGLE_CSE_ID")

        has_search_provider = any(os.getenv(key) for key in search_providers) or google_cse_configured

        if not has_search_provider:
            self.result.add_warning("No web search provider configured")
            self.result.add_info("Web search will be unavailable - consider configuring TAVILY_API_KEY")
        else:
            configured_providers = [key for key in search_providers if os.getenv(key)]
            if google_cse_configured:
                configured_providers.append("GOOGLE_CSE")
            self.result.add_info(f"🔍 Web search providers: {', '.join(configured_providers)}")

    def _validate_input_files(self):
        """Validate required input files"""
        input_dir = self.base_path / "input"

        required_files = {
            "course_syllabus.docx": self._validate_docx_file,
            "template.docx": self._validate_template_docx,
            "guidelines.md": self._validate_markdown_file
        }

        for filename, validator_func in required_files.items():
            file_path = input_dir / filename

            if not file_path.exists():
                self.result.add_error(f"Required input file missing: {file_path}")
                continue

            if not file_path.is_file():
                self.result.add_error(f"Path exists but is not a file: {file_path}")
                continue

            # Run specific validator
            try:
                validator_func(file_path)
            except Exception as e:
                self.result.add_error(f"Error validating {filename}: {str(e)}")

        self.result.add_info("📄 Input files validation completed")

    def _validate_docx_file(self, file_path: Path):
        """Validate DOCX file can be read and contains content"""
        try:
            doc = Document(str(file_path))

            # Check if document has content
            total_text = ""
            for paragraph in doc.paragraphs:
                total_text += paragraph.text

            if len(total_text.strip()) < 100:  # Minimum content check
                self.result.add_warning(f"{file_path.name} seems to have very little content")
            else:
                self.result.add_info(f"✅ {file_path.name} - {len(total_text)} characters")

            # Check for tables (optional but good to know)
            if doc.tables:
                self.result.add_info(f"📊 {file_path.name} contains {len(doc.tables)} tables")

        except Exception as e:
            self.result.add_error(f"Cannot read {file_path.name}: {str(e)}")

    def _validate_template_docx(self, file_path: Path):
        """Validate template DOCX with specific template requirements"""
        self._validate_docx_file(file_path)  # Basic DOCX validation

        try:
            doc = Document(str(file_path))
            content_text = "\n".join([p.text for p in doc.paragraphs]).lower()

            # Check for key template sections
            required_sections = [
                "discovery",
                "engagement",
                "consolidation",
                "learning objectives",
                "wlo"
            ]

            missing_sections = []
            for section in required_sections:
                if section not in content_text:
                    missing_sections.append(section)

            if missing_sections:
                self.result.add_warning(f"Template may be missing sections: {', '.join(missing_sections)}")
            else:
                self.result.add_info("✅ Template contains expected section indicators")

            # Check for time allocations
            if "85 minutes" not in content_text and "42 minutes" not in content_text:
                self.result.add_warning("Template may not contain expected time allocations")

        except Exception as e:
            self.result.add_error(f"Error validating template structure: {str(e)}")

    def _validate_markdown_file(self, file_path: Path):
        """Validate Markdown file can be read and contains content"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            if len(content.strip()) < 500:  # Minimum content for guidelines
                self.result.add_warning(f"{file_path.name} seems to have very little content")
            else:
                self.result.add_info(f"✅ {file_path.name} - {len(content)} characters")

            # Check for key guideline sections
            content_lower = content.lower()
            expected_sections = [
                "citation",
                "assessment",
                "multimedia",
                "building blocks",
                "wlo"
            ]

            found_sections = [section for section in expected_sections if section in content_lower]

            if len(found_sections) < 3:
                self.result.add_warning(f"Guidelines may be incomplete - found: {', '.join(found_sections)}")
            else:
                self.result.add_info(f"📋 Guidelines contains expected sections: {', '.join(found_sections)}")

        except UnicodeDecodeError:
            self.result.add_error(f"{file_path.name} contains invalid UTF-8 encoding")
        except Exception as e:
            self.result.add_error(f"Cannot read {file_path.name}: {str(e)}")

    def _validate_configuration_files(self):
        """Validate configuration files"""
        config_dir = self.base_path / "config"

        config_files = {
            "sections.json": self._validate_sections_json,
            "course_config.yaml": self._validate_course_config_yaml,
            "template_mapping.yaml": self._validate_yaml_file,
            "building_blocks_requirements.yaml": self._validate_yaml_file
        }

        for filename, validator_func in config_files.items():
            file_path = config_dir / filename

            if not file_path.exists():
                if filename in ["sections.json", "course_config.yaml"]:
                    self.result.add_error(f"Required config file missing: {file_path}")
                else:
                    self.result.add_warning(f"Optional config file missing: {file_path}")
                continue

            try:
                validator_func(file_path)
            except Exception as e:
                self.result.add_error(f"Error validating {filename}: {str(e)}")

        self.result.add_info("⚙️  Configuration files validation completed")

    def _validate_sections_json(self, file_path: Path):
        """Validate sections.json configuration"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        if not isinstance(data, list):
            self.result.add_error("sections.json must contain a list of sections")
            return

        if len(data) == 0:
            self.result.add_error("sections.json cannot be empty")
            return

        required_fields = ["id", "title", "description"]

        for i, section in enumerate(data):
            if not isinstance(section, dict):
                self.result.add_error(f"sections.json[{i}] must be an object")
                continue

            missing_fields = [field for field in required_fields if field not in section]
            if missing_fields:
                self.result.add_error(f"sections.json[{i}] missing fields: {', '.join(missing_fields)}")

        self.result.add_info(f"✅ sections.json - {len(data)} sections configured")

    def _validate_course_config_yaml(self, file_path: Path):
        """Validate course_config.yaml"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)

        if not isinstance(data, dict):
            self.result.add_error("course_config.yaml must be a YAML object")
            return

        # Check for expected top-level sections
        expected_sections = ["course", "learning_phases", "agents"]
        found_sections = [section for section in expected_sections if section in data]

        if len(found_sections) < 2:
            self.result.add_warning(f"course_config.yaml may be incomplete - found: {', '.join(found_sections)}")

        # Validate learning phases if present
        if "learning_phases" in data:
            phases = data["learning_phases"]
            expected_phases = ["discovery", "engagement", "consolidation"]

            for phase in expected_phases:
                if phase not in phases:
                    self.result.add_warning(f"Missing learning phase: {phase}")
                elif "duration_minutes" not in phases[phase]:
                    self.result.add_warning(f"Phase {phase} missing duration_minutes")

        self.result.add_info("✅ course_config.yaml structure validated")

    def _validate_yaml_file(self, file_path: Path):
        """Basic YAML file validation"""
        with open(file_path, 'r', encoding='utf-8') as f:
            yaml.safe_load(f)
        self.result.add_info(f"✅ {file_path.name} - valid YAML")

    def _validate_dependencies(self):
        """Validate Python dependencies and imports"""
        try:
            # Test critical imports
            import tiktoken
            self.result.add_info("✅ tiktoken available for token counting")
        except ImportError:
            self.result.add_warning("tiktoken not available - using fallback token estimation")

        try:
            import docx
            self.result.add_info("✅ python-docx available for document processing")
        except ImportError:
            self.result.add_error("python-docx not available - required for DOCX file processing")

        try:
            import yaml
            self.result.add_info("✅ PyYAML available for configuration files")
        except ImportError:
            self.result.add_error("PyYAML not available - required for configuration files")

        try:
            from langchain_openai import ChatOpenAI, AzureChatOpenAI
            self.result.add_info("✅ langchain-openai available for LLM integration")
        except ImportError:
            self.result.add_error("langchain-openai not available - required for AI agents")

        try:
            from langgraph.graph import StateGraph
            self.result.add_info("✅ langgraph available for workflow orchestration")
        except ImportError:
            self.result.add_error("langgraph not available - required for workflow execution")

        self.result.add_info("🔧 Dependencies validation completed")

    def print_results(self):
        """Print validation results in a user-friendly format"""
        print("\n" + "="*60)
        print("📋 INPUT VALIDATION RESULTS")
        print("="*60)

        if self.result.info:
            print("\n📌 Information:")
            for info in self.result.info:
                print(f"   {info}")

        if self.result.warnings:
            print("\n⚠️  Warnings:")
            for warning in self.result.warnings:
                print(f"   ⚠️  {warning}")

        if self.result.errors:
            print("\n❌ Errors:")
            for error in self.result.errors:
                print(f"   ❌ {error}")

        print("\n" + "="*60)

        if self.result.is_valid:
            print("✅ VALIDATION PASSED - Ready to generate content!")
        else:
            print("❌ VALIDATION FAILED - Please fix errors before continuing")
            print("\n💡 Next steps:")
            print("   1. Fix the errors listed above")
            print("   2. Run validation again")
            print("   3. Check the README.md for setup instructions")

        print("="*60 + "\n")


def validate_inputs(base_path: str = ".") -> bool:
    """Convenience function to run full input validation"""
    validator = InputValidator(base_path)
    result = validator.validate_all()
    validator.print_results()
    return result.is_valid